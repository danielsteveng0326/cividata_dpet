"""
Script para crear una plantilla base de Word para respuestas a derechos de petición.
Esta plantilla incluye marcadores de posición que serán reemplazados dinámicamente.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def crear_plantilla_respuesta():
    """Crea una plantilla Word profesional para respuestas"""
    
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # ============================================
    # ENCABEZADO INSTITUCIONAL
    # ============================================
    
    # Logo y nombre de la entidad (centrado)
    header_institucional = doc.add_paragraph()
    header_institucional.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_institucional.add_run('ALCALDÍA MUNICIPAL')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    
    # Subtítulo
    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheader.add_run('Gestión Documental')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    
    # Línea separadora
    doc.add_paragraph('_' * 90)
    
    # Espacio
    doc.add_paragraph()
    
    # ============================================
    # INFORMACIÓN DE LA COMUNICACIÓN
    # ============================================
    
    # Ciudad y fecha (alineado a la derecha)
    ciudad_fecha = doc.add_paragraph()
    ciudad_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = ciudad_fecha.add_run('{{CIUDAD}}, {{FECHA}}')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Información del destinatario
    destinatario = doc.add_paragraph()
    run = destinatario.add_run('Señor(a):\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = destinatario.add_run('{{NOMBRE_PETICIONARIO}}\n')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = destinatario.add_run('{{DIRECCION_PETICIONARIO}}\n')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = destinatario.add_run('Ciudad')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Asunto
    asunto = doc.add_paragraph()
    run = asunto.add_run('Asunto: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = asunto.add_run('Respuesta a Derecho de Petición - Radicado {{RADICADO}}')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Saludo
    saludo = doc.add_paragraph()
    run = saludo.add_run('Respetado(a) señor(a):')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # ============================================
    # CUERPO DE LA RESPUESTA
    # ============================================
    
    # Marcador para el cuerpo de la respuesta
    cuerpo = doc.add_paragraph()
    run = cuerpo.add_run('{{CUERPO_RESPUESTA}}')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cuerpo.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ============================================
    # DESPEDIDA Y FIRMA
    # ============================================
    
    despedida = doc.add_paragraph()
    run = despedida.add_run('Cordialmente,')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Línea de firma
    firma_linea = doc.add_paragraph()
    run = firma_linea.add_run('_' * 40)
    run.font.size = Pt(11)
    
    # Nombre del funcionario
    nombre_funcionario = doc.add_paragraph()
    run = nombre_funcionario.add_run('{{NOMBRE_FUNCIONARIO}}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # Cargo
    cargo = doc.add_paragraph()
    run = cargo.add_run('{{CARGO_FUNCIONARIO}}')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # Entidad
    entidad = doc.add_paragraph()
    run = entidad.add_run('Alcaldía Municipal')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ============================================
    # PIE DE PÁGINA
    # ============================================
    
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pie.add_run('_______________________________________________________________')
    run.font.size = Pt(8)
    
    info_pie = doc.add_paragraph()
    info_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_pie.add_run('Dirección: [Dirección de la Alcaldía] | Teléfono: [Teléfono] | Email: [Email]')
    run.font.size = Pt(8)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Guardar plantilla
    doc.save('plantillas_word/plantilla_respuesta_peticion.docx')
    print("✅ Plantilla creada exitosamente en: plantillas_word/plantilla_respuesta_peticion.docx")
    print("\nMarcadores de posición incluidos:")
    print("  - {{CIUDAD}}")
    print("  - {{FECHA}}")
    print("  - {{NOMBRE_PETICIONARIO}}")
    print("  - {{DIRECCION_PETICIONARIO}}")
    print("  - {{RADICADO}}")
    print("  - {{CUERPO_RESPUESTA}}")
    print("  - {{NOMBRE_FUNCIONARIO}}")
    print("  - {{CARGO_FUNCIONARIO}}")

if __name__ == '__main__':
    crear_plantilla_respuesta()
