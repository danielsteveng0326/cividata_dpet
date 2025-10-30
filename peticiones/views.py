# views.py - ARCHIVO COMPLETO ACTUALIZADO
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.paginator import Paginator
from django.db.models import Q
from django.http import JsonResponse, HttpResponseForbidden
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.views.generic import ListView, DetailView
from datetime import timedelta
from .models import Peticion, ProcesamientoIA
from .forms import PeticionForm
from .services.gemini_service import GeminiTranscriptionService
from .services.asistente_respuesta_service import AsistenteRespuestaService
import threading
import json
import logging

logger = logging.getLogger(__name__)


def puede_ver_peticion(user, peticion):
    """
    Verifica si un usuario puede ver una petición específica.
    - Oficina Jurídica (prefijo 111) puede ver TODAS las peticiones
    - Superuser puede ver TODAS las peticiones
    - Otras dependencias solo pueden ver sus propias peticiones
    """
    # Oficina Jurídica o superuser pueden ver todo
    if (user.dependencia and user.dependencia.prefijo == '111') or \
       (user.cedula == '1020458606' and user.is_superuser):
        return True
    
    # Otras dependencias solo ven sus propias peticiones
    return peticion.dependencia == user.dependencia


@login_required
def index(request):
    """Vista principal con estadísticas"""
    from datetime import date
    from .services.dias_habiles_service import DiasHabilesService
    
    # Filtrar peticiones según dependencia del usuario
    # Si es Jefe Jurídica (dependencia 111) o superuser, puede ver todas
    if (request.user.dependencia and request.user.dependencia.prefijo == '111') or \
       (request.user.cedula == '1020458606' and request.user.is_superuser):
        peticiones_queryset = Peticion.objects.all()
    else:
        # Solo ver peticiones de su dependencia
        peticiones_queryset = Peticion.objects.filter(dependencia=request.user.dependencia)
    
    total_peticiones = peticiones_queryset.count()
    sin_responder = peticiones_queryset.filter(estado='sin_responder').count()
    respondidas = peticiones_queryset.filter(estado='respondido').count()
    
    # Calcular peticiones próximas a vencer (menos de 3 días hábiles restantes)
    hoy = date.today()
    proximas_vencer = 0
    
    # Obtener peticiones sin responder con fecha de vencimiento
    peticiones_sin_responder = peticiones_queryset.filter(
        estado='sin_responder',
        fecha_vencimiento__isnull=False
    )
    
    for peticion in peticiones_sin_responder:
        # Calcular días hábiles restantes
        dias_restantes = DiasHabilesService.contar_dias_habiles_entre_fechas(hoy, peticion.fecha_vencimiento)
        
        # Si quedan 3 días hábiles o menos, está próxima a vencer
        if 0 <= dias_restantes <= 3:
            proximas_vencer += 1
    
    peticiones_recientes = peticiones_queryset.order_by('-fecha_radicacion')[:5]
    
    context = {
        'total_peticiones': total_peticiones,
        'sin_responder': sin_responder,
        'respondidas': respondidas,
        'proximas_vencer': proximas_vencer,
        'peticiones_recientes': peticiones_recientes,
    }
    return render(request, 'peticiones/index.html', context)


@login_required
def crear_peticion(request):
    """Vista para crear nueva petición"""
    if request.method == 'POST':
        form = PeticionForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            peticion = form.save(commit=False)
            
            # Asignar automáticamente la dependencia del usuario logueado
            peticion.dependencia = request.user.dependencia
            
            peticion.save()
            
            # Procesar PDF con IA en segundo plano
            def procesar_en_background():
                gemini_service = GeminiTranscriptionService()
                gemini_service.procesar_peticion_completa(peticion)
            
            # Ejecutar procesamiento en hilo separado
            thread = threading.Thread(target=procesar_en_background)
            thread.daemon = True
            thread.start()
            
            messages.success(
                request, 
                f'Petición creada exitosamente con radicado: {peticion.radicado}. '
                f'El documento está siendo procesado por IA.'
            )
            return redirect('detalle_peticion', radicado=peticion.radicado)
    else:
        form = PeticionForm(user=request.user)
    
    return render(request, 'peticiones/crear_peticion.html', {'form': form})


class ListaPeticiones(LoginRequiredMixin, ListView):
    """Vista para listar todas las peticiones"""
    model = Peticion
    template_name = 'peticiones/lista_peticiones.html'
    context_object_name = 'peticiones'
    paginate_by = 10
    
    def get_queryset(self):
        # Filtrar por dependencia del usuario
        # Si es Jefe Jurídica (111) o superuser, puede ver todas
        if (self.request.user.dependencia and self.request.user.dependencia.prefijo == '111') or \
           (self.request.user.cedula == '1020458606' and self.request.user.is_superuser):
            # Jefe Jurídica o superuser pueden ver todas las peticiones
            queryset = Peticion.objects.select_related('dependencia').all()
        else:
            # Solo ver peticiones de su dependencia
            queryset = Peticion.objects.select_related('dependencia').filter(dependencia=self.request.user.dependencia)
        
        # Filtros de búsqueda
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(radicado__icontains=search) |
                Q(peticionario_nombre__icontains=search) |
                Q(peticionario_id__icontains=search)
            )
        
        estado = self.request.GET.get('estado')
        if estado:
            queryset = queryset.filter(estado=estado)
        
        fuente = self.request.GET.get('fuente')
        if fuente:
            queryset = queryset.filter(fuente=fuente)
        
        return queryset.order_by('-fecha_radicacion')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['search'] = self.request.GET.get('search', '')
        context['estado_filtro'] = self.request.GET.get('estado', '')
        context['fuente_filtro'] = self.request.GET.get('fuente', '')
        return context


@login_required
def detalle_peticion(request, radicado):
    """Vista detalle de una petición específica"""
    peticion = get_object_or_404(Peticion, radicado=radicado)
    
    # Verificar permisos de acceso
    if not puede_ver_peticion(request.user, peticion):
        messages.error(request, 'No tienes permiso para ver esta petición.')
        return redirect('index')
    
    # Obtener información del procesamiento IA si existe
    try:
        procesamiento_ia = peticion.procesamiento_ia
    except ProcesamientoIA.DoesNotExist:
        procesamiento_ia = None
    
    context = {
        'peticion': peticion,
        'procesamiento_ia': procesamiento_ia,
    }
    return render(request, 'peticiones/detalle_peticion.html', context)


@login_required
@csrf_exempt
def reprocesar_peticion(request, radicado):
    """Vista AJAX para reprocesar una petición con IA"""
    if request.method == 'POST':
        peticion = get_object_or_404(Peticion, radicado=radicado)
        
        # Verificar permisos de acceso
        if not puede_ver_peticion(request.user, peticion):
            return JsonResponse({'success': False, 'message': 'No tienes permiso para esta acción'})
        
        def reprocesar_en_background():
            gemini_service = GeminiTranscriptionService()
            gemini_service.reanalizar_peticion(peticion)
        
        # Ejecutar reprocesamiento en hilo separado
        thread = threading.Thread(target=reprocesar_en_background)
        thread.daemon = True
        thread.start()
        
        return JsonResponse({
            'success': True,
            'message': 'Reprocesamiento iniciado correctamente'
        })
    
    return JsonResponse({'success': False, 'message': 'Método no permitido'})


@login_required
def cambiar_estado_peticion(request, radicado):
    """Vista para cambiar el estado de una petición con archivos adjuntos"""
    if request.method == 'POST':
        from .forms import MarcarRespondidoForm
        from django.utils import timezone
        
        peticion = get_object_or_404(Peticion, radicado=radicado)
        
        # Verificar permisos de acceso
        if not puede_ver_peticion(request.user, peticion):
            messages.error(request, 'No tienes permiso para modificar esta petición.')
            return redirect('index')
        nuevo_estado = request.POST.get('nuevo_estado')
        
        if nuevo_estado == 'respondido':
            # Usar formulario para validar archivos
            form = MarcarRespondidoForm(request.POST, request.FILES, instance=peticion)
            
            if form.is_valid():
                peticion = form.save(commit=False)
                peticion.estado = 'respondido'
                peticion.fecha_respuesta = timezone.now()
                peticion.save()
                
                messages.success(
                    request, 
                    f'Petición {radicado} marcada como respondida con archivos adjuntos correctamente.'
                )
            else:
                # Mostrar errores del formulario
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(request, f'{form.fields[field].label}: {error}')
                return redirect('detalle_peticion', radicado=radicado)
        elif nuevo_estado == 'sin_responder':
            peticion.estado = nuevo_estado
            peticion.save()
            
            messages.success(
                request, 
                f'Estado de la petición {radicado} actualizado a: {peticion.get_estado_display()}'
            )
        else:
            messages.error(request, 'Estado inválido')
    
    return redirect('detalle_peticion', radicado=radicado)


@login_required
def editar_peticionario(request, radicado):
    """Vista para editar los datos del peticionario"""
    peticion = get_object_or_404(Peticion, radicado=radicado)
    
    # Verificar permisos de acceso
    if not puede_ver_peticion(request.user, peticion):
        messages.error(request, 'No tienes permiso para editar esta petición.')
        return redirect('index')
    
    if request.method == 'POST':
        form = EditarPeticionarioForm(request.POST, instance=peticion)
        if form.is_valid():
            form.save()
            messages.success(
                request, 
                f'Datos del peticionario actualizados correctamente para {radicado}'
            )
            
            # Si es una petición AJAX, devolver JSON
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': 'Datos actualizados correctamente',
                    'peticionario_nombre': peticion.peticionario_nombre or '',
                    'peticionario_id': peticion.peticionario_id or '',
                    'peticionario_telefono': peticion.peticionario_telefono or '',
                    'peticionario_correo': peticion.peticionario_correo or '',
                    'peticionario_direccion': peticion.peticionario_direccion or ''
                })
            
            return redirect('lista_peticiones')
        else:
            # Si hay errores en formulario AJAX
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'errors': form.errors
                })
            
            messages.error(request, 'Error al actualizar los datos')
    else:
        form = EditarPeticionarioForm(instance=peticion)
    
    # Para peticiones normales (no AJAX)
    context = {
        'form': form,
        'peticion': peticion
    }
    return render(request, 'peticiones/editar_peticionario.html', context)


@login_required
@csrf_exempt
def obtener_datos_peticionario(request, radicado):
    """Vista AJAX para obtener los datos actuales del peticionario"""
    if request.method == 'GET':
        peticion = get_object_or_404(Peticion, radicado=radicado)
        
        # Verificar permisos de acceso
        if not puede_ver_peticion(request.user, peticion):
            return JsonResponse({'success': False, 'message': 'No tienes permiso para esta acción'})
        
        return JsonResponse({
            'success': True,
            'peticionario_nombre': peticion.peticionario_nombre or '',
            'peticionario_id': peticion.peticionario_id or '',
            'peticionario_telefono': peticion.peticionario_telefono or '',
            'peticionario_correo': peticion.peticionario_correo or '',
            'peticionario_direccion': peticion.peticionario_direccion or ''
        })
    
    return JsonResponse({'success': False, 'message': 'Método no permitido'})


# ========================================
# NUEVAS VISTAS PARA ASISTENTE IA
# ========================================

@login_required
@csrf_exempt
def iniciar_asistente_respuesta(request, radicado):
    """
    Inicia el proceso de asistente inteligente para generar respuesta
    """
    if request.method == 'POST':
        try:
            peticion = get_object_or_404(Peticion, radicado=radicado)
            
            # Verificar permisos de acceso
            if not puede_ver_peticion(request.user, peticion):
                return JsonResponse({'success': False, 'error': 'No tienes permiso para esta acción'})
            
            # Verificar que la petición tenga transcripción
            if not peticion.transcripcion_completa:
                return JsonResponse({
                    'success': False,
                    'error': 'La petición debe estar procesada por IA primero'
                })
            
            # Iniciar análisis con IA
            asistente_service = AsistenteRespuestaService()
            analisis = asistente_service.analizar_peticion_y_generar_preguntas(peticion)
            
            # Guardar análisis en sesión
            request.session[f'analisis_{radicado}'] = analisis
            
            return JsonResponse({
                'success': True,
                'analisis': analisis
            })
            
        except Exception as e:
            logger.error(f"Error en asistente para {radicado}: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
    
    return JsonResponse({'success': False, 'error': 'Método no permitido'})


@login_required
def mostrar_asistente_respuesta(request, radicado):
    """
    Muestra la interfaz del asistente con las preguntas generadas
    """
    peticion = get_object_or_404(Peticion, radicado=radicado)
    
    # Verificar permisos de acceso
    if not puede_ver_peticion(request.user, peticion):
        messages.error(request, 'No tienes permiso para acceder a esta petición.')
        return redirect('index')
    
    # Obtener análisis de la sesión
    analisis = request.session.get(f'analisis_{radicado}')
    
    if not analisis:
        messages.error(request, 'Debe iniciar el análisis primero')
        return redirect('detalle_peticion', radicado=radicado)
    
    context = {
        'peticion': peticion,
        'analisis': analisis
    }
    
    return render(request, 'peticiones/asistente_respuesta.html', context)


@login_required
@csrf_exempt
def procesar_respuestas_asistente(request, radicado):
    """
    Procesa las respuestas del usuario y genera la respuesta sugerida
    """
    if request.method == 'POST':
        try:
            peticion = get_object_or_404(Peticion, radicado=radicado)
            
            # Verificar permisos de acceso
            if not puede_ver_peticion(request.user, peticion):
                return JsonResponse({'success': False, 'error': 'No tienes permiso para esta acción'})
            
            # Obtener respuestas del formulario
            data = json.loads(request.body)
            respuestas_usuario = data.get('respuestas', [])
            
            if not respuestas_usuario:
                return JsonResponse({
                    'success': False,
                    'error': 'No se proporcionaron respuestas'
                })
            
            # Generar respuesta con IA
            asistente_service = AsistenteRespuestaService()
            resultado = asistente_service.generar_respuesta_sugerida(peticion, respuestas_usuario)
            
            # Evaluar calidad de la respuesta
            evaluacion = asistente_service.evaluar_calidad_respuesta(resultado['respuesta_sugerida'])
            
            return JsonResponse({
                'success': True,
                'respuesta_sugerida': resultado['respuesta_sugerida'],
                'evaluacion': evaluacion,
                'fecha_generacion': resultado['fecha_generacion']
            })
            
        except Exception as e:
            logger.error(f"Error procesando respuestas para {radicado}: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
    
    return JsonResponse({'success': False, 'error': 'Método no permitido'})


@login_required
def historial_asistente(request, radicado):
    """
    Muestra el historial de respuestas generadas por el asistente
    """
    peticion = get_object_or_404(Peticion, radicado=radicado)
    
    # Verificar permisos de acceso
    if not puede_ver_peticion(request.user, peticion):
        messages.error(request, 'No tienes permiso para acceder a esta petición.')
        return redirect('index')
    
    # Aquí podrías implementar un modelo para guardar el historial
    # Por ahora, mostraremos la última respuesta de la sesión
    
    context = {
        'peticion': peticion
    }
    
    return render(request, 'peticiones/historial_asistente.html', context)


@login_required
@csrf_exempt
def descargar_respuesta_word(request, radicado):
    """
    Descarga la respuesta generada en formato Word (.docx) usando plantilla institucional
    """
    if request.method == 'POST':
        try:
            from docx import Document
            from django.http import HttpResponse
            from datetime import datetime
            from django.conf import settings
            import io
            import os
            
            peticion = get_object_or_404(Peticion, radicado=radicado)
            
            # Verificar permisos de acceso
            if not puede_ver_peticion(request.user, peticion):
                return JsonResponse({'success': False, 'error': 'No tienes permiso para esta acción'})
            
            # Obtener el contenido de la respuesta del POST
            data = json.loads(request.body)
            contenido_respuesta = data.get('contenido_respuesta', '')
            
            if not contenido_respuesta:
                return JsonResponse({
                    'success': False,
                    'error': 'No se proporcionó contenido para la respuesta'
                })
            
            # Obtener datos del usuario autenticado
            if not request.user.is_authenticated:
                return JsonResponse({
                    'success': False,
                    'error': 'Debe iniciar sesión para generar documentos'
                })
            
            usuario = request.user
            # Obtener ciudad de la dependencia, o usar valor por defecto
            if usuario.dependencia and usuario.dependencia.ciudad:
                ciudad = usuario.dependencia.ciudad
            else:
                ciudad = 'Ciudad'
            nombre_funcionario = usuario.nombre_completo
            cargo_funcionario = usuario.cargo
            
            # Ruta de la plantilla
            plantilla_path = os.path.join(settings.BASE_DIR, 'plantillas_word', 'plantilla_respuesta_peticion.docx')
            
            # Verificar que existe la plantilla
            if not os.path.exists(plantilla_path):
                return JsonResponse({
                    'success': False,
                    'error': 'No se encontró la plantilla de Word. Por favor contacte al administrador.'
                })
            
            # Cargar la plantilla
            doc = Document(plantilla_path)
            
            # Preparar datos para reemplazo
            fecha_actual = datetime.now()
            
            # Usar diccionario de meses en español (independiente del locale del sistema)
            meses_espanol = {
                1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
                5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
                9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
            }
            fecha_formateada = f"{fecha_actual.day} de {meses_espanol[fecha_actual.month]} de {fecha_actual.year}"
            
            # Diccionario de reemplazos
            reemplazos = {
                '{{CIUDAD}}': ciudad,
                '{{FECHA}}': fecha_formateada,
                '{{NOMBRE_PETICIONARIO}}': peticion.peticionario_nombre or 'Ciudadano(a)',
                '{{DIRECCION_PETICIONARIO}}': peticion.peticionario_direccion or 'Dirección no especificada',
                '{{RADICADO}}': peticion.radicado,
                '{{CUERPO_RESPUESTA}}': contenido_respuesta,
                '{{NOMBRE_FUNCIONARIO}}': nombre_funcionario,
                '{{CARGO_FUNCIONARIO}}': cargo_funcionario
            }
            
            # Función para reemplazar texto en párrafos
            def reemplazar_en_parrafo(parrafo, reemplazos):
                for clave, valor in reemplazos.items():
                    if clave in parrafo.text:
                        # Guardar el formato del párrafo
                        inline = parrafo.runs
                        for i in range(len(inline)):
                            if clave in inline[i].text:
                                # Reemplazar el texto manteniendo el formato
                                inline[i].text = inline[i].text.replace(clave, valor)
            
            # Reemplazar en todos los párrafos del documento
            for parrafo in doc.paragraphs:
                reemplazar_en_parrafo(parrafo, reemplazos)
            
            # Reemplazar en tablas si las hay
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            reemplazar_en_parrafo(parrafo, reemplazos)
            
            # Guardar documento en memoria
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            # Preparar respuesta HTTP
            response = HttpResponse(
                file_stream.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            filename = f'Respuesta_{peticion.radicado}_{datetime.now().strftime("%Y%m%d")}.docx'
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            logger.error(f"Error generando documento Word para {radicado}: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
    
    return JsonResponse({'success': False, 'error': 'Método no permitido'})